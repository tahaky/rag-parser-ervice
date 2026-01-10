from typing import Dict, Any, List
from docx import Document
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from app.parsers.base import BaseParser


class DocxParser(BaseParser):
    """Parser for DOCX documents."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX document into structured JSON."""
        doc = Document(file_path)

        # Extract document metadata
        core_props = doc.core_properties
        metadata = self.build_metadata(
            title=core_props.title or "",
            author=core_props.author or "",
            subject=core_props.subject or "",
            created=str(core_props.created) if core_props.created else None,
            modified=str(core_props.modified) if core_props.modified else None,
        )

        # Parse sections
        sections = []
        current_section = {"level": 0, "title": "Document", "paragraphs": [], "tables": []}
        
        total_text_length = 0
        total_tables = 0
        total_images = 0

        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = element._element
                paragraph = doc.element.body.index(element)
                if paragraph < len(doc.paragraphs):
                    p = doc.paragraphs[paragraph]
                    text = p.text.strip()
                    if text:
                        para_data = {
                            "text": text,
                            "style": p.style.name if p.style else "Normal",
                            "hash": self.calculate_hash(text),
                        }
                        current_section["paragraphs"].append(para_data)
                        total_text_length += len(text)

                        # Check if it's a heading to create new section
                        if p.style and "Heading" in p.style.name:
                            if current_section["paragraphs"] or current_section["tables"]:
                                sections.append(current_section)
                            
                            level = 1
                            if "Heading" in p.style.name:
                                try:
                                    level = int(p.style.name.replace("Heading", "").strip() or "1")
                                except:
                                    level = 1
                            
                            current_section = {
                                "level": level,
                                "title": text,
                                "paragraphs": [],
                                "tables": [],
                            }

            elif isinstance(element, CT_Tbl):
                table_idx = len([e for e in doc.element.body[:doc.element.body.index(element)] if isinstance(e, CT_Tbl)])
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    table_data = self._extract_table(table)
                    current_section["tables"].append(table_data)
                    total_tables += 1

        # Add last section
        if current_section["paragraphs"] or current_section["tables"]:
            sections.append(current_section)

        # Count images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                total_images += 1

        # Build structure
        structure = {
            "format": "docx",
            "sections": sections,
        }

        # Build stats
        stats = self.build_stats(
            total_pages=len(sections),
            total_text_length=total_text_length,
            total_tables=total_tables,
            total_images=total_images,
        )

        return {
            "format": "docx",
            "metadata": metadata,
            "structure": structure,
            "stats": stats,
        }

    def _extract_table(self, table: Table) -> Dict[str, Any]:
        """Extract table data."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        table_text = " ".join(" ".join(row) for row in rows)
        return {
            "rows": rows,
            "row_count": len(rows),
            "col_count": len(rows[0]) if rows else 0,
            "hash": self.calculate_hash(table_text),
        }
