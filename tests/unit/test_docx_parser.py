import pytest
import os
import tempfile
from docx import Document
from app.parsers.docx_parser import DocxParser


class TestDocxParser:
    """Test cases for DOCX parser."""

    @pytest.fixture
    def sample_docx(self):
        """Create a sample DOCX file for testing."""
        # Create a temporary DOCX file
        fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        # Create document with content
        doc = Document()
        doc.core_properties.title = "Test Document"
        doc.core_properties.author = "Test Author"
        
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("This is the second paragraph.")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Cell 1"
        table.cell(1, 1).text = "Cell 2"
        
        doc.add_heading("Section 2", level=1)
        doc.add_paragraph("This is in section 2.")
        
        doc.save(temp_path)
        
        yield temp_path
        
        # Cleanup
        if os.path.exists(temp_path):
            os.remove(temp_path)

    def test_parse_docx(self, sample_docx):
        """Test parsing a DOCX document."""
        parser = DocxParser()
        result = parser.parse(sample_docx)

        # Check format
        assert result["format"] == "docx"

        # Check metadata
        assert "metadata" in result
        assert result["metadata"]["title"] == "Test Document"
        assert result["metadata"]["author"] == "Test Author"

        # Check structure
        assert "structure" in result
        assert result["structure"]["format"] == "docx"
        assert "sections" in result["structure"]
        sections = result["structure"]["sections"]
        assert len(sections) > 0

        # Check stats
        assert "stats" in result
        assert result["stats"]["total_text_length"] > 0
        assert result["stats"]["total_tables"] >= 1

    def test_parser_extracts_text(self, sample_docx):
        """Test that parser extracts text content."""
        parser = DocxParser()
        result = parser.parse(sample_docx)

        sections = result["structure"]["sections"]
        found_text = False
        
        for section in sections:
            if section["paragraphs"]:
                found_text = True
                # Check that paragraphs have required fields
                para = section["paragraphs"][0]
                assert "text" in para
                assert "hash" in para
                assert len(para["hash"]) == 32  # MD5 hash

        assert found_text, "Should extract some text"

    def test_parser_extracts_tables(self, sample_docx):
        """Test that parser extracts tables."""
        parser = DocxParser()
        result = parser.parse(sample_docx)

        sections = result["structure"]["sections"]
        found_table = False
        
        for section in sections:
            if section["tables"]:
                found_table = True
                table = section["tables"][0]
                assert "rows" in table
                assert "row_count" in table
                assert "col_count" in table
                assert "hash" in table
                assert table["row_count"] > 0
                assert table["col_count"] > 0

        assert found_table, "Should extract at least one table"

    def test_empty_docx(self):
        """Test parsing an empty DOCX document."""
        fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        
        try:
            doc = Document()
            doc.save(temp_path)
            
            parser = DocxParser()
            result = parser.parse(temp_path)
            
            assert result["format"] == "docx"
            assert "structure" in result
            assert "metadata" in result
            assert "stats" in result
            
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
